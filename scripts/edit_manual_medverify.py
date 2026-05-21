from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SOURCE = Path(r"C:\Users\adive\Downloads\INT428_ProjectManual (1).docx")
OUTPUT = Path(r"C:\Users\adive\Downloads\INT428_ProjectManual_MedVerify_Edited.docx")

BLUE = "1F4E79"
TEAL = "008C8C"
LIGHT_BLUE = "EAF4FB"
LIGHT_TEAL = "E8F6F5"
LIGHT_GRAY = "F4F6F8"
MID_GRAY = "D9E2EC"
TEXT = RGBColor(31, 41, 55)


def set_table_fixed_layout(table, total_width=6.6):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(int(total_width * 1440)))


def set_cell_width(cell, width_inches):
    cell.width = Inches(width_inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="D9E2EC", sz="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), sz)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            margins.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_font(run, size=None, bold=None, italic=None, color=None, name="Aptos"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc, text="", style=None, size=10.5, bold=False, italic=False, color=None, align=None, before=0, after=6):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.08
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_font(r, size=size, bold=bold, italic=italic, color=color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}" if level <= 3 else None)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_font(r, size=16 if level == 1 else 12.5, bold=True, color=BLUE if level == 1 else TEAL)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.22)
        p.paragraph_format.first_line_indent = Inches(-0.12)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.05
        r = p.add_run("- " + item)
        set_font(r, size=10.3)


def add_numbered(doc, items):
    for idx, item in enumerate(items, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.05
        r = p.add_run(f"{idx}. {item}")
        set_font(r, size=10.3)


def add_note_box(doc, title, body, fill=LIGHT_TEAL):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_font(r, size=10.5, bold=True, color=BLUE)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Inches(0.18)
    p2.paragraph_format.right_indent = Inches(0.18)
    p2.paragraph_format.space_after = Pt(8)
    p2.paragraph_format.line_spacing = 1.08
    r2 = p2.add_run(body)
    set_font(r2, size=9.8)


def add_table(doc, headers, rows, widths=None):
    for row in rows:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.04
        label = str(row[0])
        r1 = p.add_run(label + ": ")
        set_font(r1, size=9.8, bold=True, color=BLUE)
        if len(row) > 1:
            r2 = p.add_run(str(row[1]))
            set_font(r2, size=9.8)
        if len(row) > 2 and str(row[2]).strip():
            r3 = p.add_run(" | " + headers[2] + ": ")
            set_font(r3, size=9.8, bold=True, color=TEAL)
            r4 = p.add_run(str(row[2]))
            set_font(r4, size=9.8)
    add_para(doc, "", after=1)
    return None


def add_code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(line)
        set_font(r, size=8.5, name="Consolas")
    add_para(doc, "", after=1)


def add_answer(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.05
    r1 = p.add_run(label + ": ")
    set_font(r1, size=10.2, bold=True, color=BLUE)
    r2 = p.add_run(text)
    set_font(r2, size=10.2)


def configure_document(doc):
    for section in doc.sections:
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    styles["Normal"]._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    styles["Normal"].font.size = Pt(10.5)

    for name, size, color in (("Heading 1", 16, BLUE), ("Heading 2", 12.5, TEAL), ("Heading 3", 11, BLUE)):
        style = styles[name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)


def clear_body(doc):
    body = doc._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def build_document():
    # Build on a clean DOCX base. The source template carries layout state that
    # causes fixed-width tables to collapse in headless rendering.
    doc = Document()
    configure_document(doc)

    # Cover
    add_para(doc, "Project-Based Assessment Manual", size=20, bold=True, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    add_para(doc, "Med-Verify: AI Medical Fact-Checking Web Application", size=16, bold=True, color=TEAL, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_para(doc, "Healthcare domain-specific generative AI system using APIs", size=11.5, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=14)

    add_table(
        doc,
        ["Field", "Project Details"],
        [
            ["Course / Assessment", "INT428 Project-Based Assessment"],
            ["Domain", "Healthcare and medical misinformation verification"],
            ["Project Type", "Hybrid, LLM-based fact-checking assistant with evidence retrieval and OCR"],
            ["Primary AI API", "Google Gemini API"],
            ["Supporting APIs", "Mistral OCR, WHO Global Health Observatory, ODPHP MyHealthfinder"],
            ["Frontend", "React 18 with Vite and CSS Modules"],
            ["Backend", "Node.js, Express, Multer, Axios"],
            ["Repository", "https://github.com/Aditya-0x/MedTech"],
        ],
        widths=[1.8, 4.8],
    )

    add_note_box(
        doc,
        "Project summary",
        "Med-Verify verifies health claims typed by users or extracted from uploaded social-media screenshots. It combines trusted public-health sources with a constrained Gemini prompt to produce a structured verdict, confidence score, explanation, key facts, warnings, and related topics.",
    )

    add_heading(doc, "1. Introduction")
    add_para(
        doc,
        "Med-Verify is a healthcare-focused generative AI application designed to reduce the risk of misleading medical advice spreading through short posts, forwarded messages, and screenshots. Instead of acting as a diagnosis tool, it checks claims against evidence sources and explains whether the available information supports, contradicts, or cannot verify the claim.",
    )
    add_bullets(
        doc,
        [
            "Users can type a claim directly or upload an image containing a medical statement.",
            "Mistral OCR extracts text from screenshots before the claim is analyzed.",
            "WHO Global Health Observatory data and ODPHP MyHealthfinder content provide domain grounding.",
            "Gemini synthesizes the evidence into a JSON verdict displayed in the React interface.",
            "The application includes clear educational-use boundaries and encourages professional medical consultation for personal decisions.",
        ],
    )

    add_heading(doc, "2. Assessment Objective")
    add_para(doc, "The project demonstrates practical AI integration, domain-specific prompt control, evidence retrieval, and responsible presentation of generated medical explanations.")
    add_bullets(
        doc,
        [
            "Identify a real-world problem: fast-moving public medical misinformation.",
            "Use generative AI APIs in a domain-constrained manner.",
            "Integrate external healthcare data sources and OCR into a working web prototype.",
            "Present technical decisions clearly through architecture, data flow, configuration, and code evidence.",
        ],
    )

    add_heading(doc, "3. Project Description")
    add_heading(doc, "Domain-Specific AI Chatbot / Assistant", level=2)
    add_para(
        doc,
        "Med-Verify behaves as a medical fact-checking assistant rather than an open-ended casual chatbot. The system receives a claim, gathers relevant public-health context, and asks Gemini to return one of four verdicts: TRUE, FALSE, MISLEADING, or UNVERIFIED. The response is restricted to structured JSON so the user interface can consistently display a verdict card with sources and warnings.",
    )
    add_table(
        doc,
        ["Capability", "Implementation In Med-Verify"],
        [
            ["Text input", "React textarea accepts direct health claims up to 1000 characters."],
            ["Image input", "Upload, drag-and-drop, or paste screenshots; backend stores the image in memory for OCR."],
            ["OCR", "Mistral OCR endpoint extracts text; Pixtral vision chat is used as a fallback."],
            ["Evidence retrieval", "Backend searches WHO indicators and MyHealthfinder topics using claim keywords."],
            ["AI synthesis", "Gemini generates a constrained fact-check verdict and supporting explanation."],
            ["Result display", "React result card presents confidence, key facts, warnings, sources, and metadata."],
        ],
        widths=[1.7, 4.9],
    )

    add_heading(doc, "4. Scope and Requirements")
    add_heading(doc, "Mandatory Requirements", level=2)
    add_table(
        doc,
        ["Requirement", "Status", "Evidence"],
        [
            ["Clear domain definition", "Met", "Healthcare claim verification and medical misinformation checking."],
            ["Use of at least one Generative AI API", "Met", "Google Gemini API is used for the final verdict and explanation."],
            ["Domain-constrained responses", "Met", "System prompt restricts verdict labels, response JSON, and medical-safety behavior."],
            ["Working prototype", "Met", "React/Vite frontend and Express backend with /api/verify route."],
        ],
        widths=[2.0, 0.8, 3.8],
    )
    add_heading(doc, "Optional Enhancements", level=2)
    add_table(
        doc,
        ["Enhancement", "Project Status"],
        [
            ["Conversation memory", "Not used; the system is single-turn to avoid carrying unsafe medical context."],
            ["Document/image-based Q&A", "Partially implemented through screenshot upload and OCR extraction."],
            ["Multilingual support", "Not implemented; current external source queries and interface are English-focused."],
            ["Simple user interface", "Implemented with dedicated tabs for typed claims and screenshot upload."],
        ],
        widths=[2.1, 4.5],
    )

    add_heading(doc, "5. Tools and Technologies")
    add_table(
        doc,
        ["Layer", "Tools / APIs", "Purpose"],
        [
            ["Frontend", "React 18, Vite, CSS Modules", "Claim input, screenshot upload, loading state, result rendering."],
            ["Backend", "Node.js, Express, Multer, Axios", "API routing, JSON handling, in-memory image upload, external API calls."],
            ["Generative AI", "Google Gemini REST API", "Evidence synthesis and structured verdict generation."],
            ["OCR", "Mistral OCR; Pixtral fallback", "Text extraction from uploaded social-media screenshots."],
            ["Evidence APIs", "WHO GHO, ODPHP MyHealthfinder", "Authoritative public-health data and guideline context."],
            ["Deployment", "Vercel configuration, local start.bat", "Frontend build plus serverless API rewrite; local development support."],
        ],
        widths=[1.25, 2.25, 3.1],
    )

    add_heading(doc, "6. Data Collection and Domain Knowledge Preparation")
    add_para(
        doc,
        "The project does not train a model or create a private dataset. Domain knowledge is prepared through trusted external sources and carefully designed keyword extraction.",
    )
    add_bullets(
        doc,
        [
            "WHO Global Health Observatory indicators are selected from claim keywords such as diabetes, obesity, vaccines, heart disease, mental health, and tobacco.",
            "ODPHP MyHealthfinder topics are queried for preventive-care and evidence-based recommendation summaries.",
            "The system prompt tells the model to cite provided WHO/Healthfinder context when relevant and avoid treatment prescriptions.",
            "Example claims in the UI include vaccine misinformation, diabetes claims, vitamin C, exercise, coffee, and heart-health statements.",
        ],
    )
    add_note_box(
        doc,
        "Domain-control approach",
        "The LLM is not allowed to answer freely as a general-purpose chatbot. It receives curated source context and must produce one of four verdict categories in a predictable JSON schema.",
        fill=LIGHT_BLUE,
    )

    add_heading(doc, "7. Model Configuration Awareness")
    add_table(
        doc,
        ["Parameter", "Configured Value", "Reason"],
        [
            ["Temperature", "0.3", "Low randomness supports factual, cautious medical explanations."],
            ["Top-p", "Not explicitly configured", "The API default is used; safety is mainly controlled through prompt constraints and source grounding."],
            ["Max output tokens", "2048", "Allows a complete JSON response with explanation, facts, warnings, and related topics."],
            ["Model fallback chain", "Gemini 2.5 Flash Lite -> 2.5 Flash -> 2.0 Flash Lite -> 2.0 Flash -> Flash Latest", "Improves reliability if a model is unavailable, rate-limited, or not found."],
        ],
        widths=[1.5, 2.35, 2.75],
    )

    add_heading(doc, "8. Deliverables")
    add_bullets(
        doc,
        [
            "Project report/manual: this completed INT428 manual.",
            "Source code: React frontend, Express backend, and serverless API wrapper.",
            "Application access: local frontend at http://localhost:5173 and backend at http://localhost:5000 during development; Vercel-ready configuration is included.",
            "Presentation slides: to be created from this architecture, workflow, and evidence summary.",
        ],
    )

    add_heading(doc, "9. Evaluation Alignment")
    add_table(
        doc,
        ["Criterion", "How Med-Verify Addresses It"],
        [
            ["Problem identification and innovation", "Targets medical misinformation in social posts and forwards using OCR plus source-grounded AI."],
            ["Technical execution", "Combines React, Express, upload handling, OCR, health APIs, and Gemini synthesis."],
            ["Prompt engineering", "Uses a strict medical fact-checker role, verdict options, and JSON output schema."],
            ["Model configuration", "Uses low temperature and explicit max output length, with documented fallback models."],
            ["Demo effectiveness", "Supports both direct claims and screenshot-based user workflows."],
        ],
        widths=[2.15, 4.45],
    )

    add_heading(doc, "10. Academic Integrity")
    add_para(
        doc,
        "The project source code and documentation should be submitted with proper attribution for all APIs and public-health data sources. The application clearly states that it is for educational fact-checking only and is not a substitute for professional medical advice.",
    )
    add_bullets(
        doc,
        [
            "Google Gemini API is credited for evidence synthesis.",
            "Mistral OCR is credited for screenshot text extraction.",
            "WHO GHO and ODPHP MyHealthfinder are credited for authoritative health evidence.",
            "All generated medical outputs include safety boundaries and consultation guidance.",
        ],
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_para(doc, "Chatbot-Based Project Evaluation Questionnaire (INT428)", size=17, bold=True, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    add_table(
        doc,
        ["Student / Project Field", "Completed Information"],
        [
            ["Student Name", "____________________________"],
            ["Roll Number", "____________________________"],
            ["Branch & Semester", "____________________________"],
            ["Project Title", "Med-Verify: AI Medical Fact-Checking Web Application"],
            ["Guide / Faculty Name", "____________________________"],
        ],
        widths=[2.0, 4.6],
    )

    add_heading(doc, "Section A: Project Overview")
    add_answer(doc, "Q1. Type of Chatbot Developed", "[x] Hybrid. The project uses generative Gemini synthesis with retrieval from WHO/MyHealthfinder and OCR preprocessing.")
    add_answer(doc, "Q2. Platform Used for Deployment", "[x] Web Application. It includes a React frontend and Express backend API; the repository also contains Vercel deployment configuration.")
    add_answer(doc, "Q3. Deployment Link / Access Details", "Local frontend: http://localhost:5173. Local backend: http://localhost:5000. Production deployment URL can be added after Vercel deployment.")

    add_heading(doc, "Section B: Model & API Details")
    add_answer(doc, "Q4. Type of API Used", "[x] Google Gemini API for model response generation. Additional APIs: Mistral OCR, WHO GHO, and ODPHP MyHealthfinder.")
    add_answer(doc, "Q5. Model Name Used", "Primary attempt: gemini-2.5-flash-lite. Fallback models: gemini-2.5-flash, gemini-2.0-flash-lite, gemini-2.0-flash, gemini-flash-latest.")
    add_answer(doc, "Q6. Model Version", "Gemini REST API v1beta endpoint, with model fallback handling implemented in server/services/geminiAgent.js.")

    add_heading(doc, "Section C: Context & Data Handling")
    add_answer(doc, "Q7. Contextual Memory Usage", "[x] No memory (single-turn chatbot). Each request is evaluated independently to keep medical context explicit and auditable.")
    add_para(doc, "Q8. Flow of Data in the Chatbot", size=10.8, bold=True, color=BLUE, after=3)
    add_numbered(
        doc,
        [
            "User enters a medical claim or uploads a screenshot in the React interface.",
            "Frontend sends JSON claim text or multipart form-data to POST /api/verify.",
            "If an image is uploaded, Multer stores it in memory and Mistral OCR extracts text.",
            "Backend validates the claim, then fetches relevant WHO and MyHealthfinder evidence in parallel.",
            "Gemini receives the claim plus evidence context and returns structured JSON.",
            "Backend sends the verdict, sources, OCR text, and metadata back to the frontend.",
            "ResultCard displays the verdict, confidence, explanation, warnings, key facts, and sources.",
        ],
    )

    add_heading(doc, "Section D: Model Configuration & Behavior")
    add_para(doc, "Q9. Model Parameters Used", size=10.8, bold=True, color=BLUE, after=3)
    add_table(
        doc,
        ["Parameter", "Value", "Not Applicable"],
        [
            ["Temperature", "0.3", ""],
            ["Top-p / Probability", "API default; not explicitly set", "[x]"],
            ["Input Token Limit", "UI claim limit: 1000 characters; model input token limit not explicitly set", ""],
            ["Output Token Limit", "2048 maxOutputTokens", ""],
        ],
        widths=[2.0, 3.7, 0.9],
    )
    add_answer(doc, "Q10. Thinking Level & Role Assignment", "Thinking level: [x] Advanced (multi-step reasoning). Role assigned: [x] Domain Expert / medical fact-checker.")

    add_heading(doc, "Section E: Technology Stack")
    add_table(
        doc,
        ["Item", "Project Stack"],
        [
            ["Frontend", "React 18, Vite, CSS Modules"],
            ["Backend", "Node.js, Express, Multer, Axios"],
            ["Database / Vector Store", "Not used; evidence is fetched live from external APIs."],
            ["Cloud / Hosting", "Vercel-ready through vercel.json; local run scripts also provided."],
            ["Environment Variables", "GEMINI_API_KEY and MISTRAL_API_KEY are required on the backend."],
        ],
        widths=[2.1, 4.5],
    )

    add_heading(doc, "Section F: Implementation Evidence")
    add_para(doc, "Q12. API Call Evidence", size=10.8, bold=True, color=BLUE, after=3)
    add_para(doc, "The backend endpoint POST /api/verify accepts either typed claim text or image upload data, then runs OCR, evidence retrieval, and Gemini fact-checking before returning a structured JSON response.")
    add_code_block(
        doc,
        [
            "router.post('/verify', async (req, res) => {",
            "  let claimText = req.body?.claim?.trim() || '';",
            "  if (req.file) {",
            "    const ocrText = await extractTextFromImage(req.file.buffer, req.file.mimetype);",
            "    claimText = claimText ? `${ocrText}\\n\\nAdditional context: ${claimText}` : ocrText;",
            "  }",
            "  const [whoData, healthfinderData] = await Promise.all([",
            "    queryWHOData(claimText),",
            "    getHealthRecommendations(claimText)",
            "  ]);",
            "  const verdict = await factCheckClaim(claimText, whoData, healthfinderData);",
            "  res.json({ success: true, claim: claimText.slice(0, 500), verdict, sources: { who: whoData, healthfinder: healthfinderData } });",
            "});",
        ],
    )
    add_para(doc, "Gemini request configuration used by the project:", size=10.5, bold=True, color=BLUE, after=3)
    add_code_block(
        doc,
        [
            "generationConfig: {",
            "  temperature: 0.3,",
            "  maxOutputTokens: 2048",
            "}",
        ],
    )
    add_answer(doc, "Q13. Chatbot Working Interface Evidence", "The interface contains two main modes: Type a Claim and Upload Screenshot. It supports example medical claims, direct text entry, screenshot drag-and-drop/paste, OCR analysis, loading feedback, and a structured result card.")
    add_answer(doc, "Q14. GitHub Repository Link", "https://github.com/Aditya-0x/MedTech")

    add_heading(doc, "Declaration")
    add_para(doc, "I confirm that the information provided above is accurate to the best of my knowledge.")
    add_para(doc, "Student Signature: ______________________        Date: ______________________")

    # Footer with concise source attribution.
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.text = ""
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("Med-Verify INT428 Project Manual | Educational use only")
        set_font(r, size=8.5, color="64748B")

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
