from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn


SOURCE = Path(r"C:\Users\adive\Downloads\INT428_ProjectManual (1).docx")
OUTPUT = Path(r"C:\Users\adive\Downloads\INT428_ProjectManual_MedVerify_SameFormat.docx")


def set_run_text(run, text):
    run.text = text


def set_plain_text(paragraph, text):
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        run = paragraph.add_run(text)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        run.font.size = Pt(12)


def replace_exact(doc, old, new):
    for paragraph in doc.paragraphs:
        if paragraph.text == old:
            set_plain_text(paragraph, new)
            return paragraph
    raise ValueError(f"Paragraph not found: {old!r}")


def replace_startswith(doc, prefix, new):
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith(prefix):
            set_plain_text(paragraph, new)
            return paragraph
    raise ValueError(f"Paragraph prefix not found: {prefix!r}")


def check_option(doc, option_text):
    for paragraph in doc.paragraphs:
        if paragraph.text == "☐ " + option_text or paragraph.text == "☑ " + option_text:
            if paragraph.runs:
                paragraph.runs[0].text = "☑"
            else:
                set_plain_text(paragraph, "☑ " + option_text)
            return paragraph
    raise ValueError(f"Checkbox option not found: {option_text!r}")


def fill_label_run(paragraph, answer):
    # Preserve the original bold label run and replace only the answer line.
    if len(paragraph.runs) >= 2:
        paragraph.runs[1].text = " " + answer
    else:
        set_plain_text(paragraph, paragraph.text.split(":", 1)[0] + ": " + answer)


def fill_label_by_prefix(doc, prefix, answer):
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith(prefix):
            fill_label_run(paragraph, answer)
            return paragraph
    raise ValueError(f"Label prefix not found: {prefix!r}")


def main():
    doc = Document(SOURCE)

    # Manual front matter and project-specific content, keeping the same paragraph positions.
    replace_exact(doc, "Domain-Specific Generative AI Chatbot Using APIs", "Med-Verify: AI Medical Fact-Checking Using APIs")
    replace_exact(
        doc,
        "Develop a chatbot tailored to a specific domain such as healthcare, education, agriculture, finance, law, or e-commerce. The chatbot must utilize Generative AI APIs and demonstrate domain customization.",
        "Developed Med-Verify, a healthcare fact-checking assistant that verifies typed medical claims and screenshot text using Gemini, Mistral OCR, WHO GHO, and MyHealthfinder APIs.",
    )
    replace_exact(doc, "Clear domain definition", "Clear domain definition: Healthcare misinformation verification")
    replace_exact(doc, "Use of at least one Generative AI API", "Use of Generative AI API: Google Gemini")
    replace_exact(doc, "Domain-constrained and meaningful responses", "Domain-constrained verdicts using WHO and MyHealthfinder evidence")
    replace_exact(doc, "Working prototype (app or notebook-based)", "Working React and Express web prototype")
    replace_exact(doc, "Conversation memory", "Conversation memory: Not used; single-turn for safety")
    replace_exact(doc, "Document-based Q&A", "Document/image-based Q&A: Screenshot OCR supported")
    replace_exact(doc, "Multilingual support", "Multilingual support: Not implemented")
    replace_exact(doc, "Simple user interface", "Simple user interface: React web app with text and image tabs")
    replace_exact(doc, "Generative AI APIs (OpenAI, Gemini, Claude, Hugging Face)", "Generative AI API: Google Gemini")
    replace_exact(doc, "Programming languages: Python (recommended), JavaScript", "Programming language: JavaScript (React + Node.js)")
    replace_exact(doc, "Frameworks: Streamlit, Flask, FastAPI, Gradio", "Frameworks: React (Vite), Express, Vercel-ready API")
    replace_exact(
        doc,
        "Although the chatbot uses Generative AI APIs, students are expected to demonstrate domain understanding through basic data collection and preparation. This does not require building datasets or training models.",
        "Med-Verify uses trusted public-health sources and keyword-based evidence retrieval. It does not train a private model or store a medical dataset.",
    )
    replace_exact(doc, "Collecting domain FAQs from trusted websites", "WHO Global Health Observatory indicators")
    replace_exact(doc, "Identifying common user queries and intents", "ODPHP MyHealthfinder evidence topics")
    replace_exact(doc, "Referring to government portals, textbooks, standards, or manuals", "Medical misinformation examples from social media claims")
    replace_exact(doc, "Curating sample questions and expected answers", "Sample claims and expected verdict categories")
    replace_exact(doc, "Improve prompt quality", "Improve prompt quality through evidence grounding")
    replace_exact(doc, "Define system instructions accurately", "Define a medical fact-checker system role")
    replace_exact(doc, "Ensure domain-relevant and reliable responses", "Ensure cautious, source-aware medical responses")
    replace_exact(
        doc,
        "Students must explain what information sources were studied and how domain knowledge influenced prompt design.",
        "Sources studied: WHO GHO, ODPHP MyHealthfinder, Mistral OCR, and Google Gemini API behavior.",
    )
    replace_exact(doc, "Controls randomness and creativity in responses", "Configured at 0.3 to reduce randomness in medical answers")
    replace_exact(doc, "Low values produce deterministic and factual answers", "Low value chosen for factual, cautious verdicts")
    replace_exact(doc, "High values increase creativity and variation", "Higher values avoided because this is a healthcare fact-checker")
    replace_exact(doc, "Controls the probability mass of token selection", "Not explicitly configured; API default is used")
    replace_exact(doc, "Lower values restrict output to safer responses", "Safety is mainly handled through prompt constraints and sources")
    replace_exact(doc, "Higher values allow broader expression", "Output is restricted to JSON verdict format")
    replace_exact(
        doc,
        "Students should justify chosen values based on the application domain and demonstrate response differences using varied settings.",
        "The selected configuration supports stable, evidence-grounded medical fact-checking.",
    )
    replace_exact(doc, "Project Report (PDF)", "Completed Project Report / Manual (PDF)")
    replace_exact(doc, "Source Code / Notebook", "Source Code / GitHub Repository")
    replace_exact(doc, "Application Link", "Application Link / Local Demo")

    # Questionnaire header: only fill the project title, leaving personal fields blank.
    header = doc.paragraphs[87]
    if len(header.runs) > 10:
        header.runs[10].text = " Med-Verify: AI Medical Fact-Checking Web Application"

    check_option(doc, "Hybrid")
    check_option(doc, "Web Application")
    fill_label_by_prefix(doc, "Deployment URL / App Link:", "localhost:5173; backend localhost:5000")
    check_option(doc, "Google Gemini API")
    fill_label_by_prefix(doc, "Model Name:", "gemini-2.5-flash-lite; fallback models configured")
    fill_label_by_prefix(doc, "Model Version / Release:", "Gemini REST API v1beta")
    check_option(doc, "No memory (single-turn chatbot)")

    # Q8 blank lines, kept in the same blank-line positions.
    doc.paragraphs[139].runs[0].text = "1. User enters claim text or uploads screenshot in React UI."
    doc.paragraphs[140].runs[0].text = "2. Frontend sends JSON or multipart form-data to POST /api/verify."
    doc.paragraphs[141].runs[0].text = "3. Mistral OCR extracts screenshot text if an image is uploaded."
    doc.paragraphs[142].runs[0].text = "4. Backend fetches WHO GHO and MyHealthfinder evidence, then Gemini returns JSON verdict."
    set_plain_text(doc.paragraphs[143], "5. ResultCard displays verdict, confidence, explanation, warnings, facts, and sources.")

    # Existing Q9 parameter table.
    table = doc.tables[0]
    table.cell(1, 1).paragraphs[0].runs[0].text = "0.3"
    table.cell(2, 1).paragraphs[0].runs[0].text = "Default"
    table.cell(2, 2).paragraphs[0].runs[-1].text = "☑"
    table.cell(3, 1).paragraphs[0].runs[0].text = "1000"
    if len(table.cell(3, 1).paragraphs[0].runs) > 1:
        table.cell(3, 1).paragraphs[0].runs[1].text = ""
    table.cell(4, 1).paragraphs[0].runs[0].text = "2048"

    check_option(doc, "Advanced (multi-step reasoning)")
    check_option(doc, "Domain Expert")

    replace_startswith(doc, "Frontend:", "Frontend: React 18, Vite, CSS Modules")
    replace_startswith(doc, "Backend:", "Backend: Node.js, Express, Multer, Axios")
    replace_startswith(doc, "Database / Vector Store:", "Database / Vector Store: Not used; live evidence APIs")
    replace_startswith(doc, "Cloud / Hosting:", "Cloud / Hosting: Vercel-ready; local start.bat")

    replace_exact(
        doc,
        "Paste a code snippet showing the chatbot making a successful API request and receiving a response.",
        "API evidence: /api/verify runs OCR, evidence lookup, Gemini fact-checking, and returns JSON.",
    )
    replace_exact(doc, "Model/API response", "Model/API response: verdict, confidence, explanation, keyFacts, warnings, relatedTopics")
    replace_exact(doc, "User input or prompt", "User input or prompt: typed claim or OCR text extracted from uploaded screenshot")
    replace_exact(doc, "Successful request execution", "Successful request execution: response includes success, sources, and metadata")
    replace_exact(
        doc,
        "Paste a screenshot of the chatbot’s user interface while interacting with a user.",
        "UI evidence: the React interface provides Type a Claim and Upload Screenshot workflows with result display.",
    )
    replace_exact(doc, "User input", "User input: medical claim text or uploaded social-media screenshot")
    replace_exact(doc, "Chatbot response", "Chatbot response: structured fact-check verdict card with sources")
    replace_exact(doc, "Platform used (web/app/messaging)", "Platform used: web application")
    replace_startswith(doc, "Repository URL:", "Repository URL: https://github.com/Aditya-0x/MedTech")

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
